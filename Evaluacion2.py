import sqlite3 as sql
import pandas as pd 
import matplotlib.pyplot as plt
from unidecode import unidecode
from docx import Document
from docx.shared import  Pt,Cm,Mm

data = sql.connect('db_personas.db')
cursor = data.cursor()  
comandoon = '''
Select p.rut, p.nombre_completo, p.nacionalidad, s.Rol, s.Sueldo, p.fecha_ingreso, p.residencia, p.fecha_de_nacimiento, p.profesion
FROM Salarios s
INNER JOIN personas p
ON p.id_rol = s.id_salarios
'''
cursor.execute(comandoon)
data = cursor.fetchall()  
#creacion del dataframe 
columnas = [i[0] for i in cursor.description]
df = pd.DataFrame(data, columns=columnas)

print (df)

def multiples_contratos(df: pd.DataFrame, inicio:int, final:int):
    
    if inicio < 0 or final > len(df):
        print("Por favor, ingrese un rango válido")
        return
    
    for i in range (inicio, final):
        persona_selc  = df.iloc[i]
        contrato1(
            persona_selc['fecha_ingreso'],
            persona_selc['Rol'],
            persona_selc['residencia'],
            persona_selc['rut'],
            persona_selc['nombre_completo'],
            persona_selc['nacionalidad'],
            persona_selc['fech_de_nacimiento'],
            persona_selc['profesion'],
            str(persona_selc['Sueldo'])
        )

pregunta_inicio = input("¿Hola, que te gustaria hacer? \n 1: Usar el filtro \n 2: Crear multiples contratos \n 3: Ver Graficos \n Ingrese un número según su preferencia 1, 2 o 3: ")

if pregunta_inicio == '1':
    #Definicion del la funcion filtro
    def filtro(df: pd.DataFrame):    
        click = input("¿Como desea buscar a la persona?(Nombre, rut, nacionalidad, sueldo, rol): ").strip().lower()
        
        if click == "nombre":
            nombreFiltro = input("Ingrese el nombre de la persona: ").strip().lower()
            nombreSNtilde = unidecode(nombreFiltro)
            df['nombre_completo'] = df['nombre_completo'].apply(lambda x: unidecode(x).lower())
            filtrodf = df[df['nombre_completo'].str.contains(nombreSNtilde, case=False)]
            
            if not filtrodf.empty:
                print(f"Se han encontrado a las siguientes personas con el nombre {nombreFiltro}:")
                print(filtrodf)
                if len(filtrodf) > 1:
                    fila_idc = int(input("Ingrese el indice de la fila que desea seleccionar: "))
                else:
                    fila_idc = filtrodf.index[0]

                persona_selc = filtrodf.iloc[filtrodf.index.get_loc(fila_idc)]
                print("Datos de la persona seleccionada: ")
                print(persona_selc)

                contratoa = input("¿Desea generar el contrato de la persona seleccionada? (si/no): ").strip().lower()
                contraoSNtilde = unidecode(contratoa)

                if contraoSNtilde == "si":
                    contrato(persona_selc)
                else:
                    print("Gracias por usar el programa")
            else:
                print(f"No se han encontrado personas con el nombre {nombreFiltro}")
        
        elif click == "rut":
            rutfiltro = input("Ingrese el rut de la persona: ").strip()
            filtrodf = df[df['rut'] == rutfiltro]
            
            if not filtrodf.empty:
                print(f"Se ha encontrado a la persona con el rut {rutfiltro}:")
                print(filtrodf)

                fila_idc = filtrodf.index[0]
                rut_selc = filtrodf.iloc[filtrodf.index.get_loc(fila_idc)]
                print("Datos de la persona seleccionada: ")
                print(rut_selc)

                contratoa = input("¿Desea generar el contrato de la persona seleccionada? (si/no): ").strip().lower()
                contraoSNtilde = unidecode(contratoa)
                
                if contraoSNtilde == "si":
                    contrato(rut_selc)
                else:
                    print("Gracias por usar el programa")

            else:
                print(f"No se han encontrado personas con el rut {rutfiltro}")
            
        
        elif click == "nacionalidad":
            nacionalidadFiltro = input("Ingrese la nacionalidad de la persona: ").strip().lower()
            df['nacionalidad'] = df['nacionalidad'].apply(lambda x: unidecode(x).lower())
            filtrodf = df[df['nacionalidad'] == nacionalidadFiltro]
            
            if not filtrodf.empty:
                print(f"Se han encontrado a las personas con la nacionalidad {nacionalidadFiltro}:")
                print(filtrodf)
                if len(filtrodf) > 1:
                    fila_idc = int(input("Ingrese el indice de la fila que desea seleccionar: "))
                else:
                    fila_idc = filtrodf.index[0]

                nacionalida_selc = filtrodf.iloc[filtrodf.index.get_loc(fila_idc)]
                print("Datos de la persona seleccionada: ")
                print(nacionalida_selc)

                contratoa = input("¿Desea generar el contrato de la persona seleccionada? (si/no): ").strip().lower()
                contraoSNtilde = unidecode(contratoa)

                if contraoSNtilde == "si":
                    contrato(nacionalida_selc)
                else:
                    print("Gracias por usar el programa")
            else:
                print(f"No se han encontrado personas con la nacionalidad {nacionalidadFiltro}")
        
        elif click == "rol":
            rolFiltro = input("Ingrese el rol de la persona: ").strip().lower()
            df['Rol'] = df['Rol'].apply(lambda x: unidecode(x).lower())
            filtrodf = df[df['Rol'] == rolFiltro]
            
            if not filtrodf.empty:
                print(f"Se han encontrado a las personas con el rol {rolFiltro}:")
                print(filtrodf)

                if len(filtrodf) > 1:
                    fila_idc = int(input("Ingrese el indice de la fila que desea seleccionar: "))
                else:
                    fila_idc = filtrodf.index[0]

                rol_selc = filtrodf.iloc[filtrodf.index.get_loc(fila_idc)]
                print("Datos de la persona seleccionada: ")
                print(rol_selc)

                contratoa = input("¿Desea generar el contrato de la persona seleccionada? (si/no): ").strip().lower()
                contraoSNtilde = unidecode(contratoa)

                if contraoSNtilde == "si":
                    contrato(rol_selc)
                else:
                    print("Gracias por usar el programa")
            else:
                print(f"No se han encontrado personas con el rol {rolFiltro}")
        
        elif click == "sueldo":
            sueldoFiltro = int(input("Ingrese el sueldo de la persona: "))
            filtrodf = df[df['Sueldo'] == sueldoFiltro]
            
            if not filtrodf.empty:
                print(f"Se ha encontrado a la persona con el sueldo {sueldoFiltro}:")
                print(filtrodf)
                if len(filtrodf) > 1:
                    fila_idc = int(input("Ingrese el indice de la fila que desea seleccionar: "))
                else:
                    fila_idc = filtrodf.index[0]

                sueldo_selc = filtrodf.iloc[filtrodf.index.get_loc(fila_idc)]
                print("Datos de la persona seleccionada: ")
                print(sueldo_selc)

                contratoa = input("¿Desea generar el contrato de la persona seleccionada? (si/no): ").strip().lower()
                contraoSNtilde = unidecode(contratoa)

                if contraoSNtilde == "si":
                    contrato(sueldo_selc)
                else:
                    print("Gracias por usar el programa")
            else:
                print(f"No se han encontrado personas con el sueldo {sueldoFiltro}")
        else:
            print("Opción no válida")
            filtro(df)

if pregunta_inicio == '2':
    multiples_contratos(df, 1, 1)

if pregunta_inicio == '3':
    while True:
        graficos_barra = input("¿Le gustaria que le mostremos un grafico de los sueldos en base a las profesiones? (si/no): ").strip().lower()
    
        if graficos_barra == 'si':
            #Calcular el sueldo promedio por profesion
            promedio_sueldo_prof = df.groupby('profesion')['Sueldo'].mean().reset_index()

            #Crear gráfico
            plt.figure(figsize=(10, 6))
            plt.bar(promedio_sueldo_prof['profesion'], promedio_sueldo_prof['Sueldo'], color='skyblue')
            plt.xlabel('Profesion')
            plt.ylabel('Promedio Sueldo')
            plt.title('Promedio de Sueldo x Profesión')
            plt.xticks(rotation=45, ha='right')
            plt.tight_layout()
            
            #Mostrar el gráfico
            plt.show()
            break
        
        elif graficos_barra == 'no':
            print("Ok, no hay problema.")
            break
        
        else:
            print('Por favor, ingrese una respuesta válida "(si/no)"')
        
    while True:
        graficos_torta = input("¿Le gustaria que le mostremos un grafico de torta de la distribución de profesiones? (si/no): ").strip().lower()
        
        if graficos_torta == 'si':
            #Calcular la distribucion de profesiones
            distribucion_prof = df['profesion'].value_counts()
            
            #Creamos grafico de torta
            plt.figure(figsize=(4, 8))    
            plt.pie(distribucion_prof, labels=distribucion_prof.index, autopct='%1.1f%%', startangle=140)
            plt.title('Distribucion de Profesiones')
            plt.axis('equal')
            
            #Mostrar el grafico
            plt.show()
            break
        elif graficos_torta == 'no':
            print('Vale no hay problema')
            break
        else:
            print('Por favor, ingrese una respuesta válida "(si/no)"')
    
    while True:    
        graficos_conteo = input("¿Le gustaría que le mostremos un grafico de conteo de profesionales por nacionalidad? (si/no): ").strip().lower()

        if graficos_conteo == 'si':
            #Calcular el conteo de profesionales por nacionalidad
            conteo_nacionalidad = df['nacionalidad'].value_counts()
            
            #Crear grafico de conteo
            plt.figure(figsize=(10, 6))
            plt.bar(conteo_nacionalidad.index, conteo_nacionalidad, color='lightgreen')
            plt.xlabel('Nacionalidad')
            plt.ylabel('Cantidad de Profesionales')
            plt.title('Conteo de Profesionales por Nacionalidad')
            plt.xticks(rotation=45, ha='right')
            plt.tight_layout()
            
            #Mostramos Grafico
            plt.show()
            break
        elif graficos_conteo == 'no':
            print('No hay problema, sigamos...')
        else:
            print('Por favor, ingrese una respuesta válida "(si/no)"')


def contrato(perso_selc):
    fecha_ingreso = perso_selc['fecha_ingreso']
    rol = perso_selc['Rol']
    residencia = perso_selc['residencia']
    rut = perso_selc['rut']
    nombre_completo = perso_selc['nombre_completo']
    nacionalidad = perso_selc['nacionalidad']
    fecha_nac = perso_selc['fecha_de_nacimiento']
    profesion = perso_selc['profesion']
    sueldo = perso_selc['Sueldo']
    contrato1(fecha_ingreso, rol, residencia, rut, nombre_completo, nacionalidad, fecha_nac, profesion, str(sueldo))
    print(f"Se ha generado el contrato de {nombre_completo} con éxito")

def contrato1(fecha_ingreso: str, rol: str, residencia: str, rut: str, nombre_completo: str, nacionalidad: str, fecha_nac: str, profesion: str, sueldo: str) -> str:
    document = Document()

    font = document.styles['Normal'].font
    font.name = 'Arial'
    font.size = Pt(10)

    sections = document.sections
    for section in sections:
        section.page_height = Cm(29.7)  
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # encabezado de la pagina
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture("header.png", width=Cm(16)) 

    # Titulo
    title = document.add_heading('CONTRATO DE TRABAJO', level=1)
    title.alignment = 1  

    
    document.add_paragraph(
        f"En Santiago de Chile, a {fecha_ingreso}, entre {nombre_completo}, de nacionalidad {nacionalidad}, fecha de nacimiento {fecha_nac},"
        f"con domicilio en {residencia}, RUT {rut}, en adelante \"el Trabajador\", por una parte; y por la otra, la Empresa XXX, RUT 12345678-9, "
        f"representada por el Sr. XXXXXXXX, ambos con domicilio en Santiago de Chile, se ha convenido el siguiente Contrato de Trabajo: \n\n"
    )

    document.add_paragraph(
        f"PRIMERO: Por el presente contrato, el Trabajador se compromete y obliga a prestar servicios personales bajo dependencia y subordinación, "
        f"de acuerdo con las instrucciones y órdenes impartidas por la Empresa, desempeñándose en el cargo de {rol}.\n\n"
    )

    document.add_paragraph(
        f"SEGUNDO: La Empresa se compromete a pagar al Trabajador por sus servicios, una remuneración mensual de {sueldo} pesos chilenos, "
        f"en forma de sueldo base. Este pago se realizará en moneda nacional y en las fechas establecidas por la Empresa.\n\n"
    )

    document.add_paragraph(
        "TERCERO: El presente contrato tendrá una duración indefinida, comenzando el día de la firma del mismo. No obstante, "
        "cualquiera de las partes podrá poner término a este contrato de acuerdo con las causales establecidas en el Código del Trabajo.\n\n"
    )

    document.add_paragraph(
        "CUARTO: En todo lo no previsto en este contrato, se estará a lo dispuesto en el Código del Trabajo y demás leyes complementarias y "
        "reglamentarias vigentes en la República de Chile.\n\n"
    )

    document.add_paragraph(
        "Leído el presente contrato y conformes las partes, lo firman en dos ejemplares del mismo tenor y a un solo efecto, en Santiago, "
        "a la fecha de inicio mencionada anteriormente.\n\n"
    )

    # firmas
    table = document.add_table(rows=2, cols=2)

    # empleado
    cell_trabajador = table.cell(0, 0)
    cell_trabajador.text = '\n\n\n\n\n\n\n\n\n_______________________________\nFirma del Trabajador\n' + nombre_completo
    cell_trabajador_rut = table.cell(1, 0)
    cell_trabajador_rut.text = 'RUT: ' + rut

    # Representante
    cell_representante = table.cell(0, 1)
    paragraph = cell_representante.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture("firma.png", width=Cm(3))  # Adjust size as needed
    paragraph.add_run("_______________________________\nFirma del Representante de la Empresa\nNombre del Representante")

    cell_representante_rut = table.cell(1, 1)
    cell_representante_rut.text = 'RUT: 12345678-9'

    # pie de la pagina
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture("footer1.png", width=Cm(16))  

    document.save(f"{nombre_completo}_contrato.docx")
    print(f"El contrato se ha guardado como {nombre_completo}_contrato.docx")
print(df())